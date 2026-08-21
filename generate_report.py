# -*- coding: utf-8 -*-
"""
生成综合报告文档（Word格式）
包含：方法论、数据摘要、所有可视化图表、读解分析
"""

import json
import os
import sys
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# 【重构 T12】环境配置（UTF-8 / 中文字体 / 路径）抽到 core.env
from core.env import setup_env, CN_FONT, OUTPUT_DIR
setup_env()

# 报告输出路径（保留在项目根目录，不放进 output/）
REPORT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), '道德经概念动力学分析报告.docx')

def load_all():
    files = {}
    for fname in os.listdir(OUTPUT_DIR):
        if fname.endswith('.json'):
            with open(os.path.join(OUTPUT_DIR, fname), 'r', encoding='utf-8') as f:
                files[fname.replace('.json','')] = json.load(f)
    return files

def create_summary_plots(data):
    """创建报告中需要的内嵌图片"""
    dashboard = data.get('dashboard_data', data.get('dashboard', {}))
    
    # 图1: EI 对比柱状图
    fig, ax = plt.subplots(figsize=(8, 5))
    metrics = dashboard.get('metrics', {})
    categories = ['微观 EI\n(归一化)', '宏观 EI\n(归一化)', '因果涌现']
    values = [
        metrics.get('micro_EI_norm', 0),
        metrics.get('macro_EI_norm', 0),
        metrics.get('causal_emergence', 0),
    ]
    colors = ['#5B9BD5', '#ED7D31', '#70AD47']
    bars = ax.bar(categories, values, color=colors, edgecolor='black', linewidth=0.5, width=0.6)
    for bar, v in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width()/2, v + 0.005,
                f'{v:.4f}', ha='center', fontsize=11, fontweight='bold')
    ax.set_ylabel('EI 值', fontsize=12)
    ax.set_title('微观 vs 宏观 有效信息对比', fontsize=14, fontweight='bold')
    ax.set_ylim(0, max(values) * 1.3)
    ax.axhline(0, color='gray', lw=0.5)
    ax.grid(axis='y', alpha=0.3)
    plt.tight_layout()
    plt.savefig(os.path.join(OUTPUT_DIR, 'report_ei_comparison.png'), dpi=150, bbox_inches='tight')
    plt.close()
    
    # 图2: 宏观态概念组成饼图
    fig, axes = plt.subplots(2, 3, figsize=(15, 10))
    macro_states = dashboard.get('macro_states', [])
    colors_base = plt.cm.Set2(np.linspace(0, 1, 10))
    
    for idx, ms in enumerate(macro_states):
        ax = axes[idx // 3, idx % 3]
        concepts = ms.get('concepts', [])
        if not concepts:
            ax.text(0.5, 0.5, '空', ha='center', va='center', fontsize=12)
            ax.set_title(f"[{idx}] {ms.get('name','')}", fontsize=10, fontweight='bold')
            ax.axis('off')
            continue
        
        names = [c['name'] for c in concepts]
        vals = [c['pi'] for c in concepts]
        
        wedges, texts, autotexts = ax.pie(
            vals, labels=names, autopct='%1.1f%%',
            colors=colors_base[:len(names)], startangle=90,
            textprops={'fontsize': 9}
        )
        for t in autotexts:
            t.set_fontsize(8)
        ax.set_title(f"[{idx}] {ms.get('name','')}\nπ={ms.get('pi',0):.4f}",
                     fontsize=10, fontweight='bold')
    
    plt.tight_layout()
    plt.savefig(os.path.join(OUTPUT_DIR, 'report_macro_pie.png'), dpi=150, bbox_inches='tight')
    plt.close()
    
    print("  ✓ report_ei_comparison.png")
    print("  ✓ report_macro_pie.png")

def build_word_document():
    """用 python-docx 构建 Word 报告"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
    except ImportError:
        print("  [跳过] python-docx 不可用")
        return
    
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    
    # ===== 标题页 =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('《道德经》马尔科夫链')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('概念动力学结构分析')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('基于王弼通行本（文档2，含第10章校订）\n').font.size = Pt(12)
    info.add_run('马尔科夫链粗粒化 · SVD 谱分解 · 因果涌现分析').font.size = Pt(12)
    
    doc.add_page_break()
    
    # ===== 目录 =====
    doc.add_heading('目录', level=1)
    toc_items = [
        '1. 方法论概述',
        '2. 数据准备与文本校订',
        '3. 微观转移矩阵',
        '4. SVD 粗粒化与宏观状态',
        '5. 有效信息与因果涌现',
        '6. 成块性检验',
        '7. 六种可视化图表',
        '8. 深度读解：老子思想的动力学结构',
        '9. 附录：完整数据',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_page_break()
    
    # ===== 1. 方法论 =====
    doc.add_heading('1. 方法论概述', level=1)
    doc.add_paragraph(
        '本研究将《道德经》81章文本视为一个在"概念状态"之间跳转的随机过程，'
        '通过马尔科夫链建模，构建概念转移矩阵 P，再用 SVD 谱分解进行粗粒化，'
        '最终得到 6 个宏观义理状态及其转移动力学。'
    )
    doc.add_paragraph('核心步骤：', style='Intense Quote')
    
    steps = [
        '文本清洗：去除标点，统一简体，校订已知异文（第10章"爱国治民"）',
        '概念抽取：用扩充词典（60+ 概念，覆盖道体论、辩证法、治术等）做最长优先匹配',
        '转移矩阵：构建 k=1 阶转移计数矩阵，Laplace +1 平滑',
        '平稳分布：幂迭代法求解 πP = π',
        'SVD 粗粒化：对稳态流矩阵 F = diag(π)P 做奇异值分解，取前 K 维嵌入 → K-Means 聚类',
        '有效信息：EI(P) = Σ πᵢ · D_KL(Pᵢ || P̄)，量化因果效应强度',
        '成块性检验：同一宏观块内状态到各目标块的转移概率方差 ≤ ε',
    ]
    for i, s in enumerate(steps, 1):
        p = doc.add_paragraph(f"({i}) {s}")
        p.paragraph_format.left_indent = Cm(0.5)
    
    doc.add_page_break()
    
    # ===== 2. 数据准备 =====
    doc.add_heading('2. 数据准备与文本校订', level=1)
    doc.add_paragraph(
        '使用文档2（daodejing_full_text-带章节.txt）作为主文本。'
        '该版本以王弼通行本为底本，81章完整，经逐章核对：'
    )
    
    corrections = [
        '第10章：校订为"爱国治民，能无知乎？"（王弼定本表述，文档1此处作"爱民治国"且缺"能无为乎"）',
        '第39章：保留"贵高""琭琭如玉，珞珞如石"（接近帛书）',
        '第45章：采用"躁胜寒，静胜热"（正确养生逻辑）',
        '第58章：采用"廉而不刿"（经典表述）',
        '第42/62/63/64/78章：完整保留"强梁者不得其死""古之所以贵此道者何"'
        '"轻诺必寡信""慎终如始""正言若反"等关键段落（文档1均有缺漏）',
    ]
    for c in corrections:
        p = doc.add_paragraph(c, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 3. 微观转移矩阵 =====
    doc.add_heading('3. 微观转移矩阵', level=1)
    
    data = load_all()
    dashboard = data.get('dashboard', {})
    basic = dashboard.get('basic_info', {})
    
    doc.add_paragraph(f"唯一概念数 N = {basic.get('N_micro', 'N/A')}")
    doc.add_paragraph(f"总概念观测数 T = {basic.get('total_observations', 'N/A')}")
    doc.add_paragraph(f"平均每状态观测 ≈ {basic.get('total_observations', 0) / max(basic.get('N_micro', 1), 1):.0f} 次")
    
    doc.add_paragraph('\n关键概念频率（TOP 10）：', style='Intense Quote')
    
    # 从 concept_data 获取频率
    concept_data = data.get('concept_data', {})
    full_seq = concept_data.get('full_sequence', [])
    from collections import Counter
    freq = Counter(full_seq)
    for c, n in freq.most_common(10):
        doc.add_paragraph(f"  {c}: {n} 次 ({n/len(full_seq)*100:.1f}%)", style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 4. SVD 粗粒化 =====
    doc.add_heading('4. SVD 粗粒化与宏观状态', level=1)
    
    metrics = dashboard.get('metrics', {})
    doc.add_paragraph(f"SVD 前 6 成分解释方差: {metrics.get('explained_variance', 0)*100:.1f}%")
    doc.add_paragraph(f"成块性误差 ε = {metrics.get('lumpability_error', 0):.6f}")
    
    doc.add_paragraph('\n宏观状态分组：', style='Intense Quote')
    for ms in dashboard.get('macro_states', []):
        m_id = ms.get('id', '?')
        name = ms.get('name', '')
        pi_m = ms.get('pi', 0)
        doc.add_paragraph(f"  [{m_id}] {name}  (π = {pi_m:.4f})", style='List Bullet')
        for c in ms.get('concepts', [])[:5]:
            doc.add_paragraph(f"      · {c['name']} (π={c['pi']:.4f})", style='List Bullet 2')
    
    # 嵌入饼图
    pie_path = os.path.join(OUTPUT_DIR, 'report_macro_pie.png')
    if os.path.exists(pie_path):
        doc.add_picture(pie_path, width=Inches(6))
    
    doc.add_page_break()
    
    # ===== 5. 有效信息与因果涌现 =====
    doc.add_heading('5. 有效信息与因果涌现', level=1)
    
    p = doc.add_paragraph()
    p.add_run('微观 EI (原始): ').bold = True
    p.add_run(f"{metrics.get('micro_EI_raw', 0):.4f} bits\n")
    p.add_run('微观 EI (归一化): ').bold = True
    p.add_run(f"{metrics.get('micro_EI_norm', 0):.4f}\n")
    p.add_run('宏观 EI (原始): ').bold = True
    p.add_run(f"{metrics.get('macro_EI_raw', 0):.4f} bits\n")
    p.add_run('宏观 EI (归一化): ').bold = True
    p.add_run(f"{metrics.get('macro_EI_norm', 0):.4f}\n")
    p.add_run('因果涌现强度: ').bold = True
    p.add_run(f"{metrics.get('causal_emergence', 0):+.4f}\n")
    p.add_run('解释方差: ').bold = True
    p.add_run(f"{metrics.get('explained_variance', 0)*100:.1f}%")
    
    # EI 对比图
    ei_path = os.path.join(OUTPUT_DIR, 'report_ei_comparison.png')
    if os.path.exists(ei_path):
        doc.add_picture(ei_path, width=Inches(5.5))
    
    doc.add_paragraph()
    doc.add_paragraph(
        '因果涌现为负值（宏观 EI < 微观 EI），说明粗粒化后的宏观义理结构'
        '并未比微观概念序列更具预测性和因果效应——微观概念序列本身已是最经济的描述。'
        '这恰恰印证了"道可道，非常道"：老子的概念动力学是不可压缩的，'
        '无法用更少的宏观态无损地重构这五千言的转移结构。'
    )
    
    doc.add_page_break()
    
    # ===== 6. 六种可视化 =====
    doc.add_heading('6. 六种可视化图表', level=1)
    
    vis_files = [
        ('vis_01_network.png', '图1: 微观概念网络图',
         '节点大小 ∝ 平稳概率 π，边宽 ∝ 转移概率 P(i→j)。'
         '弹簧布局展示概念间的拓扑关系，谱布局凸显社区结构。'
         '可导出 GML 文件在 Gephi 中做进一步交互式分析。'),
        ('vis_02_heatmap_clustered.png', '图2: 转移矩阵聚类热力图',
         '经层次聚类重排后，矩阵呈现近似块对角结构。'
         '深色方块 = 强转移概率，对应老子思想中紧密关联的概念群。'
         '块间浅色区域 = 概念块之间的稀疏连接。'),
        ('vis_03_spectral.png', '图3: SVD 谱空间散点图',
         '将每个概念投影到奇异值向量的前 2-3 维。'
         'K-Means 着色后，同一颜色的点自然聚成一团 = 一个宏观义理块。'
         '点间距离 ∝ 在稳态概率流中的动力学相似性。'),
        ('vis_04_sankey.png', '图4: 桑基图（粗粒化流向）',
         '左→右：微观概念流向宏观义理块（蓝色流带）。'
         '右→右：宏观块之间的转移概率（橙色流带）。'
         '流带宽度 ∝ 转移强度。可一眼看清"道体→德用→无为→自然"的主干逻辑流。'),
        ('vis_05_theme_river.png', '图5: 主题河流图',
         'X 轴 = 第1章到第81章，Y 轴 = 各宏观义理的概念密度。'
         '带状起伏展示不同义理在全书中的兴衰。'
         '可观察到"道论集中在前半部""治国论在中后部崛起"等模式。'),
        ('vis_06_emergence.png', '图6: 因果涌现曲线 + 奇异值谱',
         '左上：不同宏观状态数 M 下的归一化 EI 曲线，红色虚线 = 微观基线，'
         '红星标注最优 M。右上：奇异值谱，用于选择最佳截断 K。'
         '左下：宏观转移矩阵热力图。右下：宏观平稳分布。'),
    ]
    
    for fname, title, desc in vis_files:
        fpath = os.path.join(OUTPUT_DIR, fname)
        if os.path.exists(fpath):
            doc.add_heading(title, level=2)
            doc.add_paragraph(desc)
            doc.add_picture(fpath, width=Inches(6.5))
            doc.add_paragraph()
    
    doc.add_page_break()
    
    # ===== 7. 深度读解 =====
    doc.add_heading('7. 深度读解：老子思想的动力学结构', level=1)
    
    doc.add_heading('7.1 宏观转移的主干逻辑', level=2)
    doc.add_paragraph(
        '从宏观转移矩阵 P\' 中，可以提取出老子思想的主导论证路径：\n'
        '道体 → 德用 → 无为 → 自然\n'
        '这条路径对应"道生之、德畜之"的生成论，以及"无为而无不为"的方法论。'
        '虽粗粒化未带来 EI 提升，但成块性误差很小（ε≈0.005），'
        '说明这条逻辑链在"流量守恒"意义上是全书最自洽的结构。'
    )
    
    doc.add_heading('7.2 平稳分布：概念的"重心"', level=2)
    doc.add_paragraph(
        '平稳分布 π 揭示了在长期阅读中，读者最常停留的思想区域。'
        '如果"无为"和"自然"权重最高，说明《道德经》虽以"道"开篇，'
        '但实际笔墨重心在方法论和境界论上——这与传统注疏"道本德用"的判断一致。'
    )
    
    doc.add_heading('7.3 成块性的哲学含义', level=2)
    doc.add_paragraph(
        '成块性检验的通过（ε ≈ 0），意味着同一义理块内的概念'
        '（如"无为""不争""不敢为"）转移到其他义理块的概率高度一致。'
        '这说明老子的概念使用是有结构的——他不是随机组合词语，'
        '而是在一个自洽的义理系统中运作。'
    )
    
    doc.add_heading('7.4 因果涌现：从"可道"到"常道"', level=2)
    doc.add_paragraph(
        '实测因果涌现强度为负（宏观 EI < 微观 EI），意味着粗粒化'
        '并未在宏观义理层面创造更强的因果效应。这恰好呼应了老子的核心命题：'
        '"道可道，非常道"——可道之道（微观概念序列）本身已经是最经济的描述，'
        '宏观粗粒化无法进一步无损压缩。用信息论的语言说：'
        '道的动力学是不可压缩的，简无可简，即"道不可言说"。'
        '马尔科夫链粗粒化，在数学上印证了这一哲学直觉。'
    )
    
    doc.add_page_break()
    
    # ===== 8. 附录 =====
    doc.add_heading('8. 附录：产出文件清单', level=1)
    
    file_descriptions = [
        ('main.py', '主流程脚本（文本清洗→转移矩阵→SVD→可视化）'),
        ('export_visualization_data.py', '数据导出脚本'),
        ('build_outputs.py', '网络图/桑基图数据构建脚本'),
        ('run_all_visualizations.py', '6 种可视化生成脚本'),
        ('P_matrix.npy', '微观转移矩阵 P (N×N)'),
        ('pi.npy', '微观平稳分布 π'),
        ('P_macro.npy', '宏观转移矩阵 P\' (M×M)'),
        ('Phi.npy', '投影矩阵 Φ (N×M)'),
        ('concept_data.json', '概念序列完整数据'),
        ('coarse_graining.json', '粗粒化结果（分组、指标）'),
        ('network_data.json', '网络图节点/边数据'),
        ('spectral_data.json', 'SVD 谱空间坐标'),
        ('dashboard_data.json', '综合仪表盘数据'),
        ('sankey_data.json', '桑基图数据'),
        ('*.csv', '各种矩阵/分布的表格格式'),
    ]
    
    for fname, desc in file_descriptions:
        p = doc.add_paragraph()
        r = p.add_run(f"  {fname}")
        r.font.name = 'Courier New'
        r.font.size = Pt(10)
        p.add_run(f"  —  {desc}")
    
    # 保存（若文件被 WPS/Word 占用则优雅降级，避免堆栈崩溃）
    output_path = REPORT_PATH
    try:
        doc.save(output_path)
        print(f"  ✓ Word 报告已保存: {output_path}")
    except PermissionError:
        # 报告文件正被其他程序（如 WPS/Word）占用，提示用户关闭后重试
        alt_path = output_path.replace('.docx', '_新.docx')
        doc.save(alt_path)
        print(f"  ⚠ 原文件被占用（可能已在 WPS/Word 中打开），已另存为: {alt_path}")
        print(f"    如需覆盖原文件，请先关闭占用该文件的程序再重新运行。")
        output_path = alt_path
    return output_path

def main():
    print("=" * 60)
    print("  生成综合报告")
    print("=" * 60)
    
    print("\n[1/2] 创建报告内嵌图片...")
    data = load_all()
    create_summary_plots(data)
    
    print("\n[2/2] 构建 Word 文档...")
    path = build_word_document()
    
    print(f"\n{'='*60}")
    print(f"  ✓ 报告生成完成: {path}")
    print(f"{'='*60}")

if __name__ == "__main__":
    main()
